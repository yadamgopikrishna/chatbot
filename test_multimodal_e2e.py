import io
import json
import os
from app import app
from PIL import Image
from docx import Document
from reportlab.pdfgen import canvas
import pandas as pd

def test_full_multimodal_e2e():
    client = app.test_client()

    # 1. Login with existing Oracle XE user
    print("--- 1. Testing Login ---")
    login_resp = client.post("/login", data={"email": "gopi1@gmail.com", "password": "gopi123"}, follow_redirects=True)
    assert login_resp.status_code == 200
    print("Logged in successfully, reached dashboard!")

    # 2. Upload a sample PDF
    print("\n--- 2. Testing PDF Upload & Text Extraction ---")
    pdf_buffer = io.BytesIO()
    c = canvas.Canvas(pdf_buffer)
    c.drawString(100, 750, "Quarterly Financial Analysis Report 2026")
    c.drawString(100, 700, "Revenue increased by 18 percent across enterprise accounts.")
    c.drawString(100, 650, "Operating margin reached 32 percent.")
    c.showPage()
    c.drawString(100, 750, "Page 2: Regional Operations")
    c.drawString(100, 700, "APAC region grew at 42 percent annual rate.")
    c.save()
    pdf_buffer.seek(0)

    pdf_upload = client.post(
        "/api/documents/upload",
        data={"file": (pdf_buffer, "Financial_Report_2026.pdf")},
        content_type="multipart/form-data"
    )
    assert pdf_upload.status_code == 200
    pdf_data = json.loads(pdf_upload.data)
    assert pdf_data["success"] is True
    uploaded_pdf = pdf_data["documents"][0]
    print(f"Uploaded PDF: {uploaded_pdf['filename']}, Pages: {uploaded_pdf['page_count']}, ID: {uploaded_pdf['doc_id']}")

    # 3. Upload a sample DOCX
    print("\n--- 3. Testing DOCX Upload ---")
    docx_doc = Document()
    docx_doc.add_heading("Project Alpha Overview", 0)
    docx_doc.add_paragraph("Project Alpha is focused on advanced multimodal autonomous agent systems.")
    docx_buffer = io.BytesIO()
    docx_doc.save(docx_buffer)
    docx_buffer.seek(0)

    docx_upload = client.post(
        "/api/documents/upload",
        data={"file": (docx_buffer, "Project_Alpha.docx")},
        content_type="multipart/form-data"
    )
    assert docx_upload.status_code == 200
    docx_data = json.loads(docx_upload.data)
    uploaded_docx = docx_data["documents"][0]
    print(f"Uploaded DOCX: {uploaded_docx['filename']}, ID: {uploaded_docx['doc_id']}")

    # 4. Upload a sample Spreadsheet
    print("\n--- 4. Testing Spreadsheet Upload & Analytics ---")
    df = pd.DataFrame({
        "Product": ["Laptop", "Monitor", "Keyboard", "Mouse", "Headset"],
        "Units_Sold": [450, 890, 1200, 1500, 670],
        "Revenue": [450000, 178000, 60000, 37500, 53600]
    })
    csv_buffer = io.BytesIO()
    df.to_csv(csv_buffer, index=False)
    csv_buffer.seek(0)

    csv_upload = client.post(
        "/api/documents/upload",
        data={"file": (csv_buffer, "product_sales.csv")},
        content_type="multipart/form-data"
    )
    assert csv_upload.status_code == 200
    csv_data = json.loads(csv_upload.data)
    uploaded_csv = csv_data["documents"][0]
    print(f"Uploaded CSV: {uploaded_csv['filename']}, Columns: {uploaded_csv['metadata'].get('columns')}")

    # 5. Test Document Library List
    print("\n--- 5. Testing Document Library Listing ---")
    list_resp = client.get("/api/documents/list")
    assert list_resp.status_code == 200
    docs_list = json.loads(list_resp.data)["documents"]
    print(f"Total documents in user library: {len(docs_list)}")
    assert len(docs_list) >= 3

    # 6. Test Document Comparison
    print("\n--- 6. Testing Document Comparison Endpoint ---")
    comp_resp = client.post("/api/documents/compare", json={
        "doc_id_a": uploaded_pdf["doc_id"],
        "doc_id_b": uploaded_docx["doc_id"]
    })
    assert comp_resp.status_code == 200
    comp_result = json.loads(comp_resp.data)["comparison"]
    print(f"Similarity: {comp_result['similarity_score']}%, Added count: {comp_result['added_count']}, Removed: {comp_result['removed_count']}")

    # 7. Test Image Forensics Detection
    print("\n--- 7. Testing Image Forensics Endpoint ---")
    img = Image.new("RGB", (512, 512), color=(100, 150, 200))
    img_buffer = io.BytesIO()
    img.save(img_buffer, format="PNG")
    img_buffer.seek(0)

    img_detect = client.post(
        "/api/images/detect",
        data={"image": (img_buffer, "camera_test.png")},
        content_type="multipart/form-data"
    )
    assert img_detect.status_code == 200
    forensics_report = json.loads(img_detect.data)["forensics"]
    print(f"Forensics Classification: {forensics_report['classification']}, AI Score: {forensics_report['ai_probability']}%")

    # 8. Test Multilingual Conversation Streaming
    print("\n--- 8. Testing SSE Streaming Chat with PDF Context in Telugu ---")
    stream_resp = client.post("/api/chat/stream", json={
        "message": "ఈ PDF లో ముఖ్యమైన విషయాలు ఏమిటి?",
        "attachments": [{
            "name": uploaded_pdf["filename"],
            "type": "pdf",
            "file_path": uploaded_pdf["file_path"],
            "doc_id": uploaded_pdf["doc_id"]
        }],
        "model": "gemini-2.5-flash"
    })
    assert stream_resp.status_code == 200
    print("Streaming SSE response stream initiated successfully!")

    print("\n=======================================================")
    print("END-TO-END MULTIMODAL ASSISTANT VALIDATION COMPLETE!")
    print("=======================================================")

if __name__ == "__main__":
    test_full_multimodal_e2e()
