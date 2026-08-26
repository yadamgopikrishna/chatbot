import os
import difflib
import re
import logging
from docx import Document

logger = logging.getLogger(__name__)

def extract_document_data(file_path):
    """
    Parses DOCX, TXT, MD, RTF files and extracts structured text, headings, and tables.
    Returns:
    {
        "filename": str,
        "file_type": str,
        "headings": list,
        "tables": list,
        "full_text": str,
        "paragraphs": list,
        "word_count": int
    }
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".docx":
        return _extract_docx(file_path)
    elif ext in [".txt", ".md", ".rtf", ".doc"]:
        return _extract_text_file(file_path, ext)
    else:
        # Fallback raw read
        return _extract_text_file(file_path, ext)


def _extract_docx(file_path):
    """Extracts content from a Word .docx file."""
    doc = Document(file_path)
    headings = []
    paragraphs = []
    full_text_lines = []
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        
        style_name = p.style.name.lower() if p.style else ""
        if "heading" in style_name or "title" in style_name:
            headings.append({"style": p.style.name, "text": text})
            full_text_lines.append(f"\n## {text}\n")
        else:
            paragraphs.append(text)
            full_text_lines.append(text)

    # Extract tables
    tables_data = []
    for table_idx, table in enumerate(doc.tables):
        table_rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(row_cells)
        if table_rows:
            tables_data.append(table_rows)
            # Format table as markdown in full text
            md_table = _format_table_to_md(table_rows)
            full_text_lines.append(f"\n[Table {table_idx + 1}]\n" + md_table + "\n")

    full_text = "\n".join(full_text_lines)
    word_count = len(full_text.split())

    return {
        "filename": os.path.basename(file_path),
        "file_type": "docx",
        "headings": headings,
        "tables": tables_data,
        "paragraphs": paragraphs,
        "full_text": full_text,
        "word_count": word_count
    }


def _extract_text_file(file_path, ext):
    """Reads plain text, markdown, or rtf files."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
    except Exception as e:
        logger.warning("Error reading file with utf-8: %s", e)
        with open(file_path, "r", encoding="latin-1", errors="ignore") as f:
            content = f.read()

    # Clean RTF tags if it's an RTF
    if ext == ".rtf":
        content = re.sub(r'{\\.*?}', '', content)
        content = re.sub(r'\\[a-z0-9]+ ?', '', content)
        content = re.sub(r'[{}\\]', '', content)

    lines = [line.strip() for line in content.splitlines() if line.strip()]
    headings = []
    for line in lines:
        if line.startswith("#") or (line.isupper() and len(line) < 80):
            headings.append({"style": "Heading", "text": line.lstrip("#").strip()})

    word_count = len(content.split())
    return {
        "filename": os.path.basename(file_path),
        "file_type": ext.lstrip("."),
        "headings": headings,
        "tables": [],
        "paragraphs": lines,
        "full_text": content,
        "word_count": word_count
    }


def _format_table_to_md(rows):
    """Converts 2D list of cells into markdown table."""
    if not rows:
        return ""
    header = "| " + " | ".join(rows[0]) + " |"
    separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"
    body_lines = ["| " + " | ".join(row) + " |" for row in rows[1:]]
    return "\n".join([header, separator] + body_lines)


def compare_documents(doc_a_text, doc_b_text, name_a="Document A", name_b="Document B"):
    """
    Compares two documents and produces a structured comparison report:
    - Added content
    - Removed content
    - Modified clauses
    - Numerical differences
    - Summary diff statistics
    """
    lines_a = [l.strip() for l in doc_a_text.splitlines() if l.strip()]
    lines_b = [l.strip() for l in doc_b_text.splitlines() if l.strip()]

    matcher = difflib.SequenceMatcher(None, lines_a, lines_b)
    added = []
    removed = []
    modified = []

    for tag, alo, ahi, blo, bhi in matcher.get_opcodes():
        if tag == "replace":
            modified.append({
                "from": "\n".join(lines_a[alo:ahi]),
                "to": "\n".join(lines_b[blo:bhi]),
                "location_a": f"Lines {alo+1}-{ahi}",
                "location_b": f"Lines {blo+1}-{bhi}"
            })
        elif tag == "delete":
            removed.append({
                "text": "\n".join(lines_a[alo:ahi]),
                "location": f"Lines {alo+1}-{ahi}"
            })
        elif tag == "insert":
            added.append({
                "text": "\n".join(lines_b[blo:bhi]),
                "location": f"Lines {blo+1}-{bhi}"
            })

    # Find number/metric differences
    numbers_a = set(re.findall(r'\b\d+(?:\.\d+)?%?\b', doc_a_text))
    numbers_b = set(re.findall(r'\b\d+(?:\.\d+)?%?\b', doc_b_text))
    new_metrics = list(numbers_b - numbers_a)[:10]
    removed_metrics = list(numbers_a - numbers_b)[:10]

    return {
        "document_a": name_a,
        "document_b": name_b,
        "total_lines_a": len(lines_a),
        "total_lines_b": len(lines_b),
        "similarity_score": round(matcher.ratio() * 100, 1),
        "added_count": len(added),
        "removed_count": len(removed),
        "modified_count": len(modified),
        "added_sections": added,
        "removed_sections": removed,
        "modified_sections": modified,
        "new_metrics": new_metrics,
        "removed_metrics": removed_metrics
    }
